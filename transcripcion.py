# transcripcion.py
# ─────────────────────────────────────────────
# VITACORE · Utilidades de exportación
# ─────────────────────────────────────────────

import os
import unicodedata
from datetime import datetime
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF


# ════════════════════════════════════════════
# FUENTE UNICODE PARA PDF
# ════════════════════════════════════════════

# Rutas candidatas para DejaVuSans en Windows
_DEJAVU_CANDIDATES = [
    r"C:\Windows\Fonts\DejaVuSans.ttf",
    r"C:\Windows\Fonts\arial.ttf",
    r"C:\Windows\Fonts\Arial.ttf",
    r"C:\Windows\Fonts\calibri.ttf",
    r"C:\Windows\Fonts\Calibri.ttf",
]

# Fuente local incluida en el proyecto (fallback garantizado)
_LOCAL_FONT = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                            "static", "fonts", "DejaVuSans.ttf")
_LOCAL_FONT_B = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                              "static", "fonts", "DejaVuSans-Bold.ttf")


def _encontrar_fuente() -> str | None:
    """Busca una fuente TTF Unicode disponible en el sistema."""
    # Primero buscar en el proyecto
    if os.path.exists(_LOCAL_FONT):
        return _LOCAL_FONT
    # Luego en Windows
    for ruta in _DEJAVU_CANDIDATES:
        if os.path.exists(ruta):
            return ruta
    return None


def _limpiar_texto(texto: str) -> str:
    """
    Elimina caracteres que fpdf no puede renderizar
    incluso con fuentes Unicode básicas.
    Reemplaza caracteres problemáticos por equivalentes ASCII.
    """
    # Reemplazos específicos conocidos
    reemplazos = {
        "─": "-",
        "━": "-",
        "═": "=",
        "│": "|",
        "□": "",
        "■": "*",
        "▪": "*",
        "•": "-",
        "→": "->",
        "←": "<-",
        "↑": "^",
        "↓": "v",
        "\u2019": "'",   # comilla derecha
        "\u2018": "'",   # comilla izquierda
        "\u201c": '"',   # comilla doble izquierda
        "\u201d": '"',   # comilla doble derecha
        "\u2013": "-",   # guión en
        "\u2014": "--",  # guión em
        "\u2026": "...", # elipsis
        "\u00b0": "°",   # grado (ok en latin-1)
    }
    for orig, repl in reemplazos.items():
        texto = texto.replace(orig, repl)

    # Como último recurso, eliminar todo lo que no sea imprimible
    # manteniendo acentos latinos (á é í ó ú ü ñ etc.)
    resultado = []
    for ch in texto:
        try:
            ch.encode("latin-1")
            resultado.append(ch)
        except (UnicodeEncodeError, UnicodeDecodeError):
            # Intentar normalizar (NFD separa acento del carácter base)
            normalizado = unicodedata.normalize("NFD", ch)
            for sub in normalizado:
                try:
                    sub.encode("latin-1")
                    resultado.append(sub)
                except (UnicodeEncodeError, UnicodeDecodeError):
                    pass  # descartar el carácter irrecuperable
    return "".join(resultado)


# ════════════════════════════════════════════
# EXPORTAR WORD
# ════════════════════════════════════════════

def exportar_word(texto: str, nombre_archivo: str, carpeta_destino: str) -> str:
    """
    Genera un archivo .docx con el informe radiológico.
    Retorna la ruta completa del archivo generado.
    """
    os.makedirs(carpeta_destino, exist_ok=True)
    ruta = os.path.join(carpeta_destino, nombre_archivo)
    doc  = DocxDocument()

    # Estilo base
    estilo           = doc.styles["Normal"]
    estilo.font.name = "Arial"
    estilo.font.size = Pt(11)

    # Título
    titulo           = doc.add_heading("Informe Radiológico", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.runs[0].font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)

    # Fecha
    fecha           = doc.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r               = fecha.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y  %H:%M')}")
    r.font.size     = Pt(9)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    # Contenido — Word maneja Unicode sin problemas
    for linea in texto.split("\n"):
        p = doc.add_paragraph()
        # Detectar encabezados de sección (ej: "HALLAZGOS")
        if linea.isupper() and len(linea.strip()) > 2 and len(linea.strip()) < 30:
            run = p.add_run(linea)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)
        elif linea.startswith("-" * 10):
            # Línea separadora — dibujar con borde inferior
            p.paragraph_format.space_after = Pt(2)
        else:
            run = p.add_run(linea)
            run.font.size = Pt(11)

    # Pie de página
    pie = doc.sections[0].footer.paragraphs[0]
    pie.text = "Generado por VITACORE · Historia Clínica Integral"
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if pie.runs:
        pie.runs[0].font.size = Pt(8)
        pie.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.save(ruta)
    return ruta


# ════════════════════════════════════════════
# EXPORTAR PDF
# ════════════════════════════════════════════

class _PDF(FPDF):
    """PDF con soporte Unicode via fuente TTF."""

    def __init__(self, fuente_ttf: str | None = None):
        super().__init__()
        self._fuente_custom = False
        self._nombre_fuente  = "Arial"   # fuente core por defecto

        if fuente_ttf and os.path.exists(fuente_ttf):
            try:
                # Intentar cargar bold también
                fuente_b = fuente_ttf.replace("DejaVuSans.ttf", "DejaVuSans-Bold.ttf")
                self.add_font("Unicode", "", fuente_ttf, uni=True)
                if os.path.exists(fuente_b):
                    self.add_font("Unicode", "B", fuente_b, uni=True)
                else:
                    self.add_font("Unicode", "B", fuente_ttf, uni=True)
                self._fuente_custom = True
                self._nombre_fuente  = "Unicode"
            except Exception as e:
                print(f"[PDF] No se pudo cargar fuente TTF ({e}), usando Arial")

    def header(self):
        if self._fuente_custom:
            self.set_font(self._nombre_fuente, "B", 13)
        else:
            self.set_font("Arial", "B", 13)
        self.set_text_color(15, 76, 129)
        self.cell(0, 10, "Informe Radiológico", align="C", new_x="LMARGIN", new_y="NEXT")
        if self._fuente_custom:
            self.set_font(self._nombre_fuente, "", 9)
        else:
            self.set_font("Arial", "", 9)
        self.set_text_color(100, 116, 139)
        self.cell(0, 6,
                  f"Fecha: {datetime.now().strftime('%d/%m/%Y  %H:%M')}",
                  align="R", new_x="LMARGIN", new_y="NEXT")
        self.ln(4)
        self.set_draw_color(14, 165, 255)
        self.set_line_width(0.5)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(4)

    def footer(self):
        self.set_y(-15)
        if self._fuente_custom:
            self.set_font(self._nombre_fuente, "", 8)
        else:
            self.set_font("Arial", "I", 8)
        self.set_text_color(148, 163, 184)
        self.cell(0, 10,
                  f"VITACORE - Historia Clinica Integral - Pag. {self.page_no()}",
                  align="C")


def exportar_pdf(texto: str, nombre_archivo: str, carpeta_destino: str) -> str:
    """
    Genera un archivo .pdf con el informe radiológico.
    Soporta caracteres Unicode (acentos, ñ, caracteres especiales médicos).
    Retorna la ruta completa del archivo generado.
    """
    os.makedirs(carpeta_destino, exist_ok=True)
    ruta = os.path.join(carpeta_destino, nombre_archivo)

    fuente_ttf = _encontrar_fuente()
    pdf = _PDF(fuente_ttf=fuente_ttf)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Si no hay fuente Unicode disponible, limpiar el texto
    if not pdf._fuente_custom:
        texto = _limpiar_texto(texto)

    # Renderizar línea por línea para dar estilo a encabezados
    nombre_fuente = pdf._nombre_fuente
    for linea in texto.split("\n"):
        linea_strip = linea.strip()
        if not linea_strip:
            pdf.ln(3)
            continue

        # Detectar encabezados (ej: "HALLAZGOS", "TÉCNICA")
        if linea_strip.isupper() and 2 < len(linea_strip) < 30:
            pdf.set_font(nombre_fuente, "B", 11)
            pdf.set_text_color(15, 76, 129)
            pdf.cell(0, 7, linea_strip, new_x="LMARGIN", new_y="NEXT")
            # Línea decorativa bajo el encabezado
            pdf.set_draw_color(14, 165, 255)
            pdf.set_line_width(0.3)
            pdf.line(pdf.get_x(), pdf.get_y(),
                     pdf.get_x() + 120, pdf.get_y())
            pdf.ln(2)
        elif set(linea_strip) <= {"-", "─", "="}:
            # Línea separadora — omitir, ya dibujamos la línea decorativa
            continue
        else:
            pdf.set_font(nombre_fuente, "", 11)
            pdf.set_text_color(15, 23, 42)
            pdf.multi_cell(0, 7, linea)

    pdf.output(ruta)
    return ruta


# ════════════════════════════════════════════
# UTILIDADES
# ════════════════════════════════════════════

def limpiar_uploads(carpeta: str, max_archivos: int = 50):
    """
    Elimina los archivos más antiguos si se supera el límite.
    """
    try:
        archivos = [
            os.path.join(carpeta, f)
            for f in os.listdir(carpeta)
            if os.path.isfile(os.path.join(carpeta, f))
        ]
        if len(archivos) > max_archivos:
            archivos.sort(key=os.path.getmtime)
            for viejo in archivos[:len(archivos) - max_archivos]:
                os.remove(viejo)
    except Exception:
        pass