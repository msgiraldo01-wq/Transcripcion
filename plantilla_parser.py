# plantilla_parser.py
# ─────────────────────────────────────────────
# VITACORE · Parser de Plantillas Radiológicas
#
# Extrae los 5 bloques clínicos de un .docx:
#   PACIENTE | ESTUDIO | TÉCNICA | HALLAZGOS | CONCLUSIÓN
#
# Estructura esperada de la plantilla:
#   - Cabecera: Paciente, Documento, Procedimiento, Fecha
#   - Sección TECNICA
#   - Sección HALLAZGOS
#   - Sección CONCLUSION
# ─────────────────────────────────────────────

from __future__ import annotations
import io
from docx import Document


# Palabras clave que identifican cada sección
_MARCADORES = {
    "tecnica":    {"tecnica", "técnica", "technique"},
    "hallazgos":  {"hallazgos", "findings", "hallazgo"},
    "conclusion": {"conclusion", "conclusión", "conclusiones",
                   "impresion", "impresión", "diagnostico", "diagnóstico"},
}

# Palabras clave de la cabecera
_CABECERA = {
    "paciente":      {"paciente", "patient", "nombre"},
    "documento":     {"documento", "doc", "cedula", "cédula", "cc", "id"},
    "procedimiento": {"procedimiento", "procedure", "estudio", "examen"},
    "fecha":         {"fecha", "date"},
    "entidad":       {"entidad", "empresa", "aseguradora", "eps"},
    "prefactura":    {"prefactura", "factura", "nro"},
}


def _es_marcador(texto: str, claves: set) -> bool:
    """Verifica si el texto es un encabezado de sección."""
    t = texto.strip().lower().rstrip(":")
    return t in claves


def _valor_cabecera(texto: str, clave: str) -> str:
    """Extrae el valor de una línea 'Clave: Valor'."""
    partes = texto.split(":", 1)
    return partes[1].strip() if len(partes) > 1 else ""


def parsear_plantilla_bytes(contenido_bytes: bytes) -> dict:
    """
    Recibe el contenido binario del .docx y retorna los 5 bloques.

    Returns:
        {
            "paciente":   str,
            "estudio":    str,
            "tecnica":    str,
            "hallazgos":  str,
            "conclusion": str,
            "ok":         bool,
            "error":      str | None,
        }
    """
    try:
        doc = Document(io.BytesIO(contenido_bytes))
        return _extraer_bloques(doc)
    except Exception as e:
        return {
            "paciente": "", "estudio": "", "tecnica": "",
            "hallazgos": "", "conclusion": "",
            "ok": False, "error": str(e),
        }


def parsear_plantilla_ruta(ruta: str) -> dict:
    """Versión que lee desde ruta de archivo."""
    try:
        doc = Document(ruta)
        return _extraer_bloques(doc)
    except Exception as e:
        return {
            "paciente": "", "estudio": "", "tecnica": "",
            "hallazgos": "", "conclusion": "",
            "ok": False, "error": str(e),
        }


def _extraer_bloques(doc: Document) -> dict:
    """Lógica central de extracción."""

    # Recolectar párrafos no vacíos
    parrafos = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # ── Estado del parser ─────────────────────────────────────
    seccion_actual = None   # "tecnica" | "hallazgos" | "conclusion" | None
    cabecera = {
        "paciente": [], "procedimiento": [], "fecha": [],
        "entidad": [], "documento": [],
    }
    bloques = {"tecnica": [], "hallazgos": [], "conclusion": []}

    # Párrafos que son firma (al final, después de conclusion)
    en_firma = False
    _firma_markers = {"atentamente", "dr", "dra", "médico", "medico",
                      "radiólogo", "radiologo", "rm", "firma"}

    for parrafo in parrafos:
        pt = parrafo.strip()
        pl = pt.lower().rstrip(":")

        # ── Detectar inicio de firma — atentamente es marcador directo ──
        if pl in {"atentamente", "atentamente,"} :
            en_firma = True
        palabras = set(pl.split())
        if len(palabras) <= 3 and palabras & {"dr", "dra", "médico", "medico", "radiologo", "radiólogo"}:
            en_firma = True

        if en_firma:
            continue  # ignorar firma

        # ── Detectar marcadores de sección ───────────────────
        if _es_marcador(pt, _MARCADORES["tecnica"]):
            seccion_actual = "tecnica"
            continue

        if _es_marcador(pt, _MARCADORES["hallazgos"]):
            seccion_actual = "hallazgos"
            continue

        if _es_marcador(pt, _MARCADORES["conclusion"]):
            seccion_actual = "conclusion"
            continue

        # ── Si estamos en una sección, acumular texto ─────────
        if seccion_actual:
            bloques[seccion_actual].append(pt)
            continue

        # ── Si no hay sección activa, es cabecera ─────────────
        pl_lower = pt.lower()

        if any(k in pl_lower for k in _CABECERA["paciente"]) and ":" in pt:
            # Solo si la línea empieza con "Paciente" (no Entidad)
            if pl_lower.startswith("paciente"):
                v = _valor_cabecera(pt, "paciente")
                if v and len(v.strip()) > 1:
                    cabecera["paciente"].append(v)

        elif any(k in pl_lower for k in _CABECERA["documento"]) and ":" in pt:
            v = _valor_cabecera(pt, "documento")
            if v and v.strip().upper() not in ("CC", "NIT", ""):
                cabecera["documento"].append(v)

        elif any(k in pl_lower for k in _CABECERA["procedimiento"]) and ":" in pt:
            v = _valor_cabecera(pt, "procedimiento")
            if v:
                cabecera["procedimiento"].append(v)

        elif any(k in pl_lower for k in _CABECERA["fecha"]) and ":" in pt:
            v = _valor_cabecera(pt, "fecha")
            if v:
                cabecera["fecha"].append(v)

        elif any(k in pl_lower for k in _CABECERA["entidad"]) and ":" in pt:
            v = _valor_cabecera(pt, "entidad")
            if v:
                cabecera["entidad"].append(v)

    # ── Ensamblar bloque PACIENTE ─────────────────────────────
    partes_pac = []
    if cabecera["paciente"]:
        partes_pac.append(f"Paciente: {' '.join(cabecera['paciente'])}")
    if cabecera["documento"]:
        partes_pac.append(f"Documento: {' '.join(cabecera['documento'])}")
    if cabecera["entidad"]:
        partes_pac.append(f"Entidad: {' '.join(cabecera['entidad'])}")
    bloque_paciente = "\n".join(partes_pac)

    # ── Ensamblar bloque ESTUDIO ──────────────────────────────
    partes_est = []
    if cabecera["procedimiento"]:
        partes_est.append(f"Estudio: {' '.join(cabecera['procedimiento'])}")
    if cabecera["fecha"]:
        partes_est.append(f"Fecha: {' '.join(cabecera['fecha'])}")
    bloque_estudio = "\n".join(partes_est)

    return {
        "paciente":   bloque_paciente,
        "estudio":    bloque_estudio,
        "tecnica":    "\n".join(bloques["tecnica"]),
        "hallazgos":  "\n".join(bloques["hallazgos"]),
        "conclusion": "\n".join(bloques["conclusion"]),
        "ok":         True,
        "error":      None,
    }


# ── Test rápido ───────────────────────────────────────────────────
if __name__ == "__main__":
    import sys
    ruta = sys.argv[1] if len(sys.argv) > 1 else "/mnt/user-data/uploads/Plantilla_1.docx"
    resultado = parsear_plantilla_ruta(ruta)
    for campo, valor in resultado.items():
        if campo not in ("ok", "error"):
            print(f"\n{'='*40}")
            print(f"[{campo.upper()}]")
            print(valor or "(vacío)")
    print(f"\nOK: {resultado['ok']}")
    if resultado['error']:
        print(f"Error: {resultado['error']}")